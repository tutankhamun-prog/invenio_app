import os
import sys
import subprocess
import tempfile
import re
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from PIL import Image, ImageTk, ImageDraw, ImageFont
from ttkbootstrap import Style
from ttkbootstrap.widgets import Button, Label, Entry, Frame
from docx import Document
import fitz  # PyMuPDF
from datetime import datetime
import json
import pytesseract  # OCR avec Tesseract

# ---------- Config Tesseract ----------
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\ofriha\Desktop\projet\tesseract-5.5.1\tesseract.exe"

# ---------- Gestion historique des mots-clés ----------
HISTORIQUE_FICHIER = "historique_recherches.json"

def charger_historique():
    if os.path.exists(HISTORIQUE_FICHIER):
        try:
            with open(HISTORIQUE_FICHIER, "r", encoding="utf-8") as f:
                mots = json.load(f)
                if isinstance(mots, list):
                    return mots
        except Exception as e:
            print(f"Erreur lecture historique : {e}")
    return []

def sauvegarder_historique(mots):
    try:
        with open(HISTORIQUE_FICHIER, "w", encoding="utf-8") as f:
            json.dump(mots, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Erreur sauvegarde historique : {e}")

historique_mots = charger_historique()

# ---------- Gestion du chemin des ressources pour l'exe ----------
def resource_path(relative_path):
    """Retourne le chemin absolu correct de la ressource, que l'on soit dans l'exe ou dans le code Python."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# ----------- Lecture des fichiers -----------
def lire_docx(path):
    try:
        doc = Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Erreur lecture DOCX {path}: {e}")
        return ""

def lire_pdf(path):
    try:
        with fitz.open(path) as doc:
            return "\n".join([page.get_text() for page in doc])
    except Exception as e:
        print(f"Erreur lecture PDF {path}: {e}")
        return ""

def lire_image_ocr(path):
    try:
        image = Image.open(path)
        texte = pytesseract.image_to_string(image, lang='fra')
        return texte
    except Exception as e:
        print(f"Erreur OCR {path}: {e}")
        return ""

def texte_en_image(texte, largeur=200, hauteur=240, max_chars_per_line=40, max_lines=10, font_path="arial.ttf"):
    try:
        image = Image.new("RGB", (largeur, hauteur), color="white")
        draw = ImageDraw.Draw(image)

        lignes = []
        for ligne in texte.strip().split("\n"):
            while len(ligne) > max_chars_per_line:
                lignes.append(ligne[:max_chars_per_line])
                ligne = ligne[max_chars_per_line:]
            lignes.append(ligne)
        lignes = lignes[:max_lines]

        texte_reduit = "\n".join(lignes)
        taille_police = 14
        try:
            font = ImageFont.truetype(font_path, taille_police)
        except:
            font = ImageFont.load_default()

        while True:
            bbox = draw.multiline_textbbox((0, 0), texte_reduit, font=font, spacing=4)
            largeur_texte = bbox[2] - bbox[0]
            hauteur_texte = bbox[3] - bbox[1]

            if largeur_texte <= (largeur - 20) and hauteur_texte <= (hauteur - 20):
                break
            taille_police -= 1
            if taille_police < 8:
                break
            try:
                font = ImageFont.truetype(font_path, taille_police)
            except:
                font = ImageFont.load_default()

        x = (largeur - largeur_texte) // 2
        y = (hauteur - hauteur_texte) // 2
        draw.multiline_text((x, y), texte_reduit, fill="black", font=font, spacing=4)
        return image
    except Exception as e:
        print(f"Erreur création image texte: {e}")
        return None

def creer_aperçu_image(path):
    try:
        if path.lower().endswith(".pdf"):
            doc = fitz.open(path)
            if len(doc) > 0:
                pix = doc[0].get_pixmap()
                return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        elif path.lower().endswith(".docx"):
            return texte_en_image(lire_docx(path), largeur=200, hauteur=240)
        elif path.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff")):
            return texte_en_image(lire_image_ocr(path), largeur=200, hauteur=240)
    except Exception as e:
        print(f"Erreur création aperçu {path}: {e}")
        return None

# ----------- Recherche -----------
def rechercher_documents(dossier, mots_cles):
    resultats = []
    try:
        for root, _, files in os.walk(dossier):
            for file in files:
                chemin_complet = os.path.join(root, file)
                try:
                    if file.endswith(".docx"):
                        texte = lire_docx(chemin_complet)
                    elif file.endswith(".pdf"):
                        texte = lire_pdf(chemin_complet)
                    elif file.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff")):
                        texte = lire_image_ocr(chemin_complet)
                    else:
                        continue
                    
                    # Vérifier si tous les mots-clés sont présents dans le texte
                    texte_lower = texte.lower()
                    tous_mots_trouves = all(mot.lower() in texte_lower for mot in mots_cles)
                    
                    if tous_mots_trouves:
                        date_modif = os.path.getmtime(chemin_complet)
                        resultats.append((root, file, date_modif))
                except Exception as e:
                    print(f"Erreur traitement fichier {chemin_complet}: {e}")
                    continue
    except Exception as e:
        print(f"Erreur parcours dossier {dossier}: {e}")
        messagebox.showerror("Erreur", f"Impossible de parcourir le dossier: {str(e)}")
    
    return resultats

def parcourir_dossier():
    dossier = filedialog.askdirectory()
    if dossier:  # Vérifie qu'un dossier a été sélectionné
        champ_dossier.delete(0, tk.END)
        champ_dossier.insert(0, dossier)

# ----------- Highlight DOCX & PDF -----------
def highlight_word_docx(input_path, output_path, mots_cles):
    try:
        doc = Document(input_path)
        for mot_cle in mots_cles:
            pattern = re.compile(re.escape(mot_cle), re.IGNORECASE)
            for para in doc.paragraphs:
                inline = para.runs
                new_runs = []
                for run in inline:
                    text = run.text
                    matches = list(pattern.finditer(text))
                    if not matches:
                        new_runs.append(run)
                        continue
                    last_end = 0
                    for m in matches:
                        if m.start() > last_end:
                            new_run = para.add_run(text[last_end:m.start()])
                            new_runs.append(new_run)
                        new_run = para.add_run(text[m.start():m.end()])
                        new_run.font.highlight_color = 7
                        new_runs.append(new_run)
                        last_end = m.end()
                    if last_end < len(text):
                        new_run = para.add_run(text[last_end:])
                        new_runs.append(new_run)
                    run.text = ""
                for r in inline:
                    p = r._element
                    p.getparent().remove(p)
                for r in new_runs:
                    para._p.append(r._element)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Erreur highlight DOCX {input_path}: {e}")
        return False

def highlight_word_pdf(input_path, output_path, mots_cles):
    try:
        doc = fitz.open(input_path)
        for page in doc:
            for mot_cle in mots_cles:
                text_instances = page.search_for(mot_cle, flags=fitz.TEXT_DEHYPHENATE)
                for inst in text_instances:
                    page.add_highlight_annot(inst)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Erreur highlight PDF {input_path}: {e}")
        return False

# ----------- Ouverture fichiers -----------
def ouvrir_fichier(path, mots_cles=None):
    try:
        if not os.path.exists(path):
            messagebox.showerror("Erreur", f"Le fichier n'existe plus: {os.path.basename(path)}")
            return
            
        suffix = os.path.splitext(path)[1].lower()
        if mots_cles:
            tmp_dir = tempfile.gettempdir()
            basename = os.path.basename(path)
            tmp_path = os.path.join(tmp_dir, f"tmp_highlight_{basename}")
            
            success = False
            if suffix == ".docx":
                success = highlight_word_docx(path, tmp_path, mots_cles)
            elif suffix == ".pdf":
                success = highlight_word_pdf(path, tmp_path, mots_cles)
            else:
                tmp_path = path
                success = True
                
            if not success:
                messagebox.showwarning("Avertissement", f"Impossible de surligner les mots-clés dans {basename}. Ouverture du fichier original.")
                tmp_path = path
        else:
            tmp_path = path

        if sys.platform.startswith('darwin'):
            subprocess.call(('open', tmp_path))
        elif os.name == 'nt':
            os.startfile(tmp_path)
        elif os.name == 'posix':
            subprocess.call(('xdg-open', tmp_path))
    except Exception as e:
        messagebox.showerror("Erreur", f"Impossible d'ouvrir le fichier: {str(e)}")

# ----------- Affichage résultats -----------
def lancer_recherche():
    global resultats_recherche, historique_mots
    dossier = champ_dossier.get().strip()
    mot_cle_input = champ_mot_cle.get().strip()
    
    # Validation des entrées
    if not dossier:
        messagebox.showerror("Erreur", "Veuillez sélectionner un dossier.")
        return
        
    if not os.path.isdir(dossier):
        messagebox.showerror("Erreur", "Le dossier spécifié n'existe pas.")
        return
        
    if not mot_cle_input:
        messagebox.showerror("Erreur", "Veuillez entrer un mot-clé.")
        return
    
    # Séparation des mots-clés par "_"
    mots_cles = [mot.strip() for mot in mot_cle_input.split('_') if mot.strip()]
    
    if not mots_cles:
        messagebox.showerror("Erreur", "Veuillez entrer au moins un mot-clé valide.")
        return
    
    # Afficher "Recherche en cours..." centré
    for widget in cadre_resultats.winfo_children():
        widget.destroy()
    
    # Créer un cadre pour centrer le message
    cadre_centre = Frame(cadre_resultats, style="TFrame")
    cadre_centre.pack(expand=True, fill="both")
    
    recherche_label = Label(cadre_centre, text="🔍 Recherche en cours...", 
                           font=("Poppins", 12), style="Result.TLabel")
    recherche_label.pack(expand=True)
    
    fenetre.update_idletasks()
    
    try:
        # Recherche des documents
        resultats_recherche = rechercher_documents(dossier, mots_cles)
        
        # Mise à jour de l'historique
        if mot_cle_input and mot_cle_input not in historique_mots:
            historique_mots.append(mot_cle_input)
            sauvegarder_historique(historique_mots)
            champ_mot_cle['values'] = historique_mots
            
        # Affichage des résultats
        if not resultats_recherche:
            messagebox.showinfo("Information", f"Aucun document trouvé contenant tous les mots-clés: {', '.join(mots_cles)}")
            for widget in cadre_resultats.winfo_children():
                widget.destroy()
        else:
            trier_afficher_resultats()
            
    except Exception as e:
        messagebox.showerror("Erreur", f"Une erreur s'est produite lors de la recherche: {str(e)}")
        for widget in cadre_resultats.winfo_children():
            widget.destroy()
        Label(cadre_resultats, text="Erreur lors de la recherche", style="Result.TLabel").pack(pady=100)

def trier_afficher_resultats(*args):
    if not resultats_recherche:
        return
        
    ordre = choix_tri.get()
    inverser = (ordre == "Plus récent → Plus ancien")
    resultats_tries = sorted(resultats_recherche, key=lambda x: x[2], reverse=inverser)
    afficher_resultats(resultats_tries)

def afficher_resultats(resultats):
    for widget in cadre_resultats.winfo_children():
        widget.destroy()
        
    if not resultats:
        Label(cadre_resultats, text="Aucun résultat à afficher", style="Result.TLabel").pack(pady=50)
        return
        
    # Récupérer les mots-clés pour le surlignage
    mots_cles = [mot.strip() for mot in champ_mot_cle.get().split('_') if mot.strip()]
        
    for dossier, fichier, date_modif in resultats:
        chemin = os.path.join(dossier, fichier)
        date_str = datetime.fromtimestamp(date_modif).strftime("%d/%m/%Y %H:%M")
        largeur_cadre = 230 if fichier.lower().endswith(".pdf") else 220
        cadre = Frame(
            cadre_resultats,
            width=largeur_cadre,
            height=400,
            padding=15,
            style="Result.TFrame",
            borderwidth=2,
            relief="solid"
        )
        cadre.pack(side="left", padx=12, pady=15)
        cadre.pack_propagate(False)
        
        try:
            img = creer_aperçu_image(chemin)
            if img:
                img = img.resize((200, 240), Image.LANCZOS)
                photo = ImageTk.PhotoImage(img)
                lbl_img = Label(cadre, image=photo, style="Result.TLabel")
                lbl_img.image = photo
                lbl_img.pack()
        except Exception as e:
            print(f"Erreur affichage aperçu {chemin}: {e}")
            Label(cadre, text="Aperçu indisponible", style="Result.TLabel").pack(pady=80)
            
        Label(
            cadre,
            text=fichier,
            font=("Poppins", 11, "bold"),
            wraplength=largeur_cadre - 40,
            justify="center",
            style="Result.TLabel"
        ).pack(pady=5)
        Label(
            cadre,
            text=f"📅 {date_str}",
            font=("Roboto", 9),
            style="Result.TLabel"
        ).pack(side="bottom", pady=5)
        
        # Bind des événements pour ouvrir le fichier
        cadre.bind("<Double-Button-1>", lambda e, p=chemin, m=mots_cles: ouvrir_fichier(p, m))
        try:
            lbl_img.bind("<Double-Button-1>", lambda e, p=chemin, m=mots_cles: ouvrir_fichier(p, m))
        except:
            pass

# ----------- Défilement horizontal -----------
def _on_mousewheel(event):
    delta = int(-1*(event.delta/120))
    canvas.xview_scroll(delta, "units")

def _on_mousewheel_linux(event):
    if event.num == 4:
        canvas.xview_scroll(-1, "units")
    elif event.num == 5:
        canvas.xview_scroll(1, "units")

# ----------- Fenêtre principale -----------
style = Style("darkly")
fenetre = style.master
fenetre.title("🔍 Recherche de mots-clés")
fenetre.geometry("1280x750")

# -- Icône --
try:
    fenetre.iconbitmap(resource_path("icone.ico"))
except Exception:
    print("Icône introuvable ou incompatible, utilisation de l'icône par défaut.")

# --- Titre INVENIO ---
img_title_path = resource_path("invenioo.png")
try:
    image_title = Image.open(img_title_path)
    photo_title = ImageTk.PhotoImage(image_title)
    Label(fenetre, image=photo_title, background=fenetre.cget("background")).pack(pady=(30, 20))
except Exception as e:
    print(f"⚠️ Impossible de charger l'image INVENIO : {e}")

# --- Formulaire ---
cadre_form = Frame(fenetre, style="Form.TFrame")
cadre_form.pack(pady=5, padx=30, fill="x")

Label(cadre_form, text="📁 Dossier :", font=("Roboto", 10, "bold"), style="TLabel").grid(row=0, column=0, sticky="e", padx=8, pady=6)
champ_dossier = Entry(cadre_form, width=70, style="TEntry", font=("Roboto", 10))
champ_dossier.grid(row=0, column=1, padx=8, pady=6, sticky="ew")
Button(cadre_form, text="Parcourir", command=parcourir_dossier, style="TButton", bootstyle="info").grid(row=0, column=2, padx=8, pady=6)

Label(cadre_form, text="🔑 Mot-clé(s) :", font=("Roboto", 10, "bold"), style="TLabel").grid(row=1, column=0, sticky="e", padx=8, pady=6)
champ_mot_cle = ttk.Combobox(cadre_form, width=70, style="TCombobox", font=("Roboto", 10))
champ_mot_cle['values'] = historique_mots
champ_mot_cle.grid(row=1, column=1, padx=8, pady=6, sticky="ew")

# Ajout d'un label d'information pour le séparateur
Label(cadre_form, text="Utiliser séparateur: '_' pour rechercher plusieurs mots à la fois (ex: mot1_mot2)", font=("Roboto", 8), style="TLabel").grid(row=2, column=1, sticky="w", padx=8)

choix_tri = ttk.Combobox(cadre_form, values=["Plus récent → Plus ancien", "Plus ancien → Plus récent"], state="readonly", style="TCombobox", font=("Roboto", 10))
choix_tri.grid(row=1, column=2, padx=8, pady=6)
choix_tri.current(0)
choix_tri.bind("<<ComboboxSelected>>", trier_afficher_resultats)

cadre_form.columnconfigure(1, weight=1)
Button(fenetre, text="🔍 Lancer la recherche", command=lancer_recherche, width=38, style="TButton", bootstyle="success").pack(pady=(15, 20))

# --- Canvas pour résultats ---
canvas = tk.Canvas(fenetre, height=400, highlightthickness=0)
scrollbar_x = tk.Scrollbar(fenetre, orient="horizontal", command=canvas.xview)
scrollable_frame = Frame(canvas, style="TFrame")
scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(xscrollcommand=scrollbar_x.set)

canvas.pack(fill="both", expand=True, padx=20)
scrollbar_x.pack(fill="x")

cadre_resultats = scrollable_frame
resultats_recherche = []

if fenetre.tk.call('tk', 'windowingsystem') == 'win32':
    canvas.bind_all("<MouseWheel>", _on_mousewheel)
elif fenetre.tk.call('tk', 'windowingsystem') == 'x11':
    canvas.bind_all("<Button-4>", _on_mousewheel_linux)
    canvas.bind_all("<Button-5>", _on_mousewheel_linux)
else:
    canvas.bind_all("<MouseWheel>", _on_mousewheel)
    
fenetre.mainloop()